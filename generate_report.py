from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Title Page ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Lab 3 / Project 2\n').bold = True
p.add_run('Autonomous Vehicle: Putting it All Together\n\n')
p.add_run('COMPENG 4DS4 - Winter 2026\n')
p.add_run('McMaster University\n\n')
p.add_run('Group 19\n\n')
p.add_run('Ali Raza - 400338538\n')
p.add_run('Zuhaib Quraishi - 400306494\n')
p.add_run('Khushi Akbari - 400320925\n')

doc.add_page_break()

# --- Declaration of Contributions ---
doc.add_heading('Declaration of Contributions', 1)

p = doc.add_paragraph()
p.add_run('Ali Raza: ').bold = True
p.add_run('Experiment 1 (Build and program PX4). Experiment 2 (hello_world application on NuttX). Experiment 3A (uORB messaging, subscribe and publish).')

p = doc.add_paragraph()
p.add_run('Zuhaib Quraishi: ').bold = True
p.add_run('Experiment 3C (motor control via test_motor). Experiment 4 (MAVLink setup and communication). Project 2 Part 1 (RC controller motor and servo control).')

p = doc.add_paragraph()
p.add_run('Khushi Akbari: ').bold = True
p.add_run('Experiment 5 (ultrasonic sensor on Raspberry Pi). Experiment 6 (camera direction detection). Project 2 Part 2 (autonomous control via MAVLink).')

doc.add_page_break()

# --- Screenshots ---
img_dir = '/tmp/report2_extracted/word/media/'

screenshots = [
    ('Experiment 2B: Hello World Application',       'image1.png'),
    ('Experiment 3A: uORB Top',                      'image4.png'),
    ('Experiment 3A: Listener sensor_combined',      'image5.png'),
    ('Experiment 3A: Listener input_rc',             'image6.png'),
    ('Experiment 3C: Motor Speed Control',           'image8.png'),
    ('Experiment 4A: MAVLink Message Stream',        'image2.png'),
    ('Experiment 4A: ATTITUDE Messages',             'image7.png'),
    ('Experiment 5: Ultrasonic Distance Readings',   'image9.png'),
    ('Experiment 6: Camera Direction Detection',     'image3.png'),
]

for title_text, img_file in screenshots:
    doc.add_heading(title_text, 2)
    img_path = os.path.join(img_dir, img_file)
    doc.add_picture(img_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

doc.save('/home/ds/Documents/Group-19/report.docx')
print('Done: report.docx')
